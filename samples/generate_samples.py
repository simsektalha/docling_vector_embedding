import os
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def ensure_dir(path: str) -> None:
    Path(path).mkdir(parents=True, exist_ok=True)


def generate_pdf(path: str) -> None:
    c = canvas.Canvas(path, pagesize=letter)
    text = c.beginText(72, 720)
    text.textLines(
        [
            "Docling RAG Starter Sample PDF",
            "",
            "This sample describes configuration for embeddings and vector DB.",
            "Set embeddings.provider to openai and model to text-embedding-3-large.",
            "For local, use huggingface all-MiniLM-L6-v2.",
            "Qdrant collection is named 'documents'.",
        ]
    )
    c.drawText(text)
    c.showPage()
    c.save()


def generate_docx(path: str) -> None:
    doc = Document()
    doc.add_heading("Product Brief", level=1)
    doc.add_paragraph("Overview: This brief describes key features and risks.")
    doc.add_heading("Risks", level=2)
    doc.add_paragraph("- Vendor lock-in if only OpenAI is used.")
    doc.add_paragraph("- Operational overhead if running vector DB ourselves.")
    doc.save(path)


def main() -> None:
    docs_dir = Path(__file__).parent / "docs"
    ensure_dir(str(docs_dir))
    generate_pdf(str(docs_dir / "sample_policy.pdf"))
    generate_docx(str(docs_dir / "product_brief.docx"))
    print(f"Generated samples in {docs_dir}")


if __name__ == "__main__":
    main()


